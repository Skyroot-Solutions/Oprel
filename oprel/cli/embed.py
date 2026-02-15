"""
Embedding commands for Oprel CLI.
"""
import argparse
import json
import sys
from pathlib import Path
from oprel.utils.logging import get_logger

logger = get_logger(__name__)

def cmd_embed(args: argparse.Namespace) -> int:
    """Generate text embeddings (similar to Ollama)"""
    from oprel.client_api import Client
    
    client = Client()
    
    # Collect texts to embed
    texts = []
    file_metadata = []  # Track file sources with metadata
    
    # Process --files flag (PDF, docs, txt, json)
    if args.files:
        for file_path in args.files:
            path = Path(file_path)
            if not path.exists():
                logger.error(f"File not found: {file_path}")
                return 1
            
            try:
                # Determine file type and extract text
                file_ext = path.suffix.lower()
                
                if file_ext == '.pdf':
                    # Extract text from PDF
                    try:
                        import PyPDF2
                        with open(path, 'rb') as f:
                            pdf_reader = PyPDF2.PdfReader(f)
                            text_content = ""
                            for page in pdf_reader.pages:
                                text_content += page.extract_text() + "\n"
                        texts.append(text_content.strip())
                        file_metadata.append({"file": str(path), "type": "pdf", "pages": len(pdf_reader.pages)})
                    except ImportError:
                        logger.error("PyPDF2 not installed. Install with: pip install PyPDF2")
                        return 1
                
                elif file_ext in ['.txt', '.md', '.py', '.js', '.json', '.yaml', '.yml']:
                    # Read text files
                    with open(path, 'r', encoding='utf-8') as f:
                        content = f.read()
                    
                    if file_ext == '.json':
                        # For JSON files, optionally parse and embed each object
                        try:
                            data = json.loads(content)
                            if isinstance(data, list):
                                # Embed each item in the list
                                for item in data:
                                    texts.append(json.dumps(item) if isinstance(item, dict) else str(item))
                                    file_metadata.append({"file": str(path), "type": "json_item"})
                            else:
                                texts.append(content)
                                file_metadata.append({"file": str(path), "type": "json"})
                        except json.JSONDecodeError:
                            texts.append(content)
                            file_metadata.append({"file": str(path), "type": "json"})
                    else:
                        texts.append(content)
                        file_metadata.append({"file": str(path), "type": file_ext[1:]})
                
                elif file_ext in ['.doc', '.docx']:
                    # Extract text from Word documents
                    try:
                        import docx
                        doc = docx.Document(path)
                        text_content = "\n".join([para.text for para in doc.paragraphs])
                        texts.append(text_content)
                        file_metadata.append({"file": str(path), "type": "docx", "paragraphs": len(doc.paragraphs)})
                    except ImportError:
                        logger.error("python-docx not installed. Install with: pip install python-docx")
                        return 1
                
                else:
                    logger.warning(f"Unsupported file type: {file_ext}, treating as text")
                    with open(path, 'r', encoding='utf-8', errors='ignore') as f:
                        texts.append(f.read())
                    file_metadata.append({"file": str(path), "type": "unknown"})
                    
            except Exception as e:
                logger.error(f"Error reading {file_path}: {e}")
                return 1

    # Process --batch flag (file with one text per line)
    if args.batch:
        path = Path(args.batch)
        if not path.exists():
            logger.error(f"Batch file not found: {path}")
            return 1
        with open(path, 'r', encoding='utf-8') as f:
            lines = [line.strip() for line in f if line.strip()]
            texts.extend(lines)
            for _ in lines:
                file_metadata.append({"file": str(path), "type": "batch_line"})

    # Process prompt argument (single text)
    if args.prompt:
        texts.append(args.prompt)
        file_metadata.append({"type": "cli_prompt"})
    
    if not texts:
        logger.error("No text to embed. Provide 'prompt', --files, or --batch.")
        return 1
        
    logger.info(f"Generating embeddings for {len(texts)} item(s)...")
    
    try:
        results = []
        import time
        start_time = time.time()
        
        # Determine model
        model_id = args.model
        
        # Generate embeddings
        # If we have many texts, we should batch them (client handles this usually?)
        # For CLI simplicity, we loop or pass list if supported. 
        # oprel.client_api.Client.embed usually takes a single text or list?
        # Checking implementation... assuming it takes list or we loop.
        
        # For now, loop
        for i, text in enumerate(texts):
            # Normalize whitespace
            if not getattr(args, 'no_normalize', False):
                text = " ".join(text.split())
                
            embedding = client.embed(
                texts=text,
                model=model_id
            )
            
            result = {
                "object": "embedding",
                "index": i,
                "embedding": embedding,
                "metadata": file_metadata[i] if i < len(file_metadata) else {}
            }
            
            if not getattr(args, 'no_texts', False):
                result["text"] = text
                
            results.append(result)
            
            if len(texts) > 1:
                print(f"Processed {i+1}/{len(texts)}", end='\r')
                
        elapsed = time.time() - start_time
        if len(texts) > 1:
            print() # Newline
            
        logger.info(f"Generated {len(texts)} embeddings in {elapsed:.2f}s")
        
        # Output
        output_format = args.format
        output_file = args.output
        
        # If files were processed and no output specified, default to embeddings.json
        if args.files and not output_file and output_format == "simple":
            output_file = "embeddings.json"
            output_format = "json"
            
        if output_file:
            with open(output_file, 'w', encoding='utf-8') as f:
                if output_format == "json":
                    json.dump(results, f, indent=2)
                elif output_format == "jsonl":
                    for r in results:
                        f.write(json.dump(r) + "\n")
                else:
                    # Simple format suitable for vector DBs?
                    json.dump(results, f) 
            print(f"Saved to {output_file}")
            return 0
            
        # Print to stdout
        if output_format == "json":
            print(json.dumps(results, indent=2))
        elif output_format == "jsonl":
            for r in results:
                print(json.dumps(r))
        else:
            # Simple output (default): Show summary to avoid flooding terminal
            for r in results:
                vec = r["embedding"]
                dim = len(vec)
                preview = f"[{', '.join(f'{x:.4f}' for x in vec[:5])}, ...]"
                print(f"Embedding {r['index']} (dim={dim}): {preview}")
            
        return 0
        
    except Exception as e:
        logger.error(f"Embedding failed: {e}")
        return 1
